#!/usr/bin/env python3
"""문서 폴더를 읽어 조문 단위로 분할한다.

사용법:
    python3 scripts/ingest.py <문서폴더> <작업폴더>

출력:
    <작업폴더>/units.jsonl  조문 단위 텍스트
    <작업폴더>/docs.json    문서 목록 (메타데이터는 비어 있음, 이후 채운다)

지원 형식: .docx .pdf .txt .md .csv
.hwp/.hwpx/.doc 는 지원하지 않는다. docx 또는 pdf 로 변환한 뒤 넣는다.
"""
import json
import re
import sys
from pathlib import Path

SUPPORTED = {".docx", ".pdf", ".txt", ".md", ".csv"}

# 조문 머리 패턴: 제1조, 제 1 조(목적), 1., 1.1, 가., ① 등
ARTICLE_RE = re.compile(
    r"^\s*(?:"
    r"제\s*\d+\s*조(?:의\s*\d+)?"          # 제12조, 제12조의2
    r"|\d+(?:\.\d+)*\.?\s"                  # 1. / 3.2.1
    r"|[가-힣]\.\s"                         # 가. 나. 다.
    r"|[①-⑮]"                              # 항 번호
    r"|\[별표\s*\d*\]|\[별지\s*\d*\]"
    r")"
)
HEAD_RE = re.compile(
    r"^\s*(제\s*\d+\s*조(?:의\s*\d+)?|\d+(?:\.\d+)*|[가-힣]\.|[①-⑮]|\[별표\s*\d*\]|\[별지\s*\d*\])"
)
ART_ONLY_RE = re.compile(r"^\s*제\s*\d+\s*조(?:의\s*\d+)?")
SUB_RE = re.compile(r"^\s*[①-⑮]")


def read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    out = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            out.append(t)
    for i, table in enumerate(doc.tables):
        out.append(f"[표{i + 1}]")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                out.append(" | ".join(cells))
    return "\n".join(out)


def read_pdf(path: Path) -> str:
    import pdfplumber

    out = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t.strip():
                out.append(t)
            for tbl in page.extract_tables() or []:
                for row in tbl:
                    cells = [(c or "").strip().replace("\n", " ") for c in row]
                    if any(cells):
                        out.append(" | ".join(cells))
    return "\n".join(out)


def read_plain(path: Path) -> str:
    for enc in ("utf-8", "cp949", "euc-kr"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def load(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return read_docx(path)
    if ext == ".pdf":
        return read_pdf(path)
    return read_plain(path)


def split_units(text: str):
    """조문 머리를 기준으로 덩어리를 나눈다. 머리가 없으면 빈 줄 기준으로 나눈다."""
    lines = [l.rstrip() for l in text.splitlines()]
    has_article = any(ARTICLE_RE.match(l) for l in lines)
    units, buf, head = [], [], None

    def flush():
        if buf and any(b.strip() for b in buf):
            units.append((head, "\n".join(buf).strip()))

    if has_article:
        last_article = None
        for line in lines:
            if ARTICLE_RE.match(line):
                flush()
                buf = [line]
                m = HEAD_RE.match(line)
                tok = m.group(1).strip() if m else line[:20]
                if ART_ONLY_RE.match(line):
                    last_article = tok
                    head = tok
                elif SUB_RE.match(line) and last_article:
                    head = f"{last_article} {tok}"
                else:
                    head = tok
            else:
                buf.append(line)
        flush()
    else:
        for line in lines:
            if not line.strip():
                flush()
                buf, head = [], None
            else:
                buf.append(line)
        flush()

    # 너무 긴 덩어리는 문장 단위로 재분할 (판정 정확도를 위해)
    result = []
    for h, body in units:
        if len(body) <= 1200:
            result.append((h, body))
            continue
        parts, cur = [], ""
        for sent in re.split(r"(?<=[.。다\.])\s+", body):
            if len(cur) + len(sent) > 1000 and cur:
                parts.append(cur)
                cur = sent
            else:
                cur = (cur + " " + sent).strip()
        if cur:
            parts.append(cur)
        for i, p in enumerate(parts):
            result.append((f"{h} ({i + 1}/{len(parts)})" if h else None, p))
    return result


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    src, work = Path(sys.argv[1]), Path(sys.argv[2])
    work.mkdir(parents=True, exist_ok=True)

    files = sorted(p for p in src.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED)
    skipped = [p.name for p in src.rglob("*") if p.is_file() and p.suffix.lower() not in SUPPORTED]

    docs, n_units = [], 0
    with (work / "units.jsonl").open("w", encoding="utf-8") as f:
        for i, path in enumerate(files, 1):
            doc_id = f"D{i:03d}"
            try:
                text = load(path)
            except Exception as e:  # noqa: BLE001
                print(f"  ! 읽기 실패 {path.name}: {e}")
                continue
            units = split_units(text)
            docs.append(
                {
                    "doc_id": doc_id,
                    "filename": path.name,
                    "path": str(path),
                    "title": path.stem,
                    "chars": len(text),
                    "units": len(units),
                    "tier": None,
                    "effective_from": None,
                    "effective_to": None,
                    "scope_default": None,
                }
            )
            for j, (head, body) in enumerate(units, 1):
                f.write(
                    json.dumps(
                        {
                            "unit_id": f"{doc_id}#u{j:03d}",
                            "doc_id": doc_id,
                            "filename": path.name,
                            "locator": head,
                            "text": body,
                        },
                        ensure_ascii=False,
                    )
                    + "\n"
                )
                n_units += 1

    (work / "docs.json").write_text(
        json.dumps(docs, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"문서 {len(docs)}건, 조문 단위 {n_units}건 → {work}/units.jsonl")
    if skipped:
        print(f"미지원 형식 {len(skipped)}건 건너뜀: {', '.join(skipped[:5])}")
        print("  .hwp/.hwpx/.doc 는 docx 또는 pdf 로 변환 후 다시 넣으세요.")


if __name__ == "__main__":
    main()